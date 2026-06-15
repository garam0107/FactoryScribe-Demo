from docx import Document


def parse_docx(path: str) -> list[dict]:
    doc = Document(path)
    blocks = []

    para_texts = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            para_texts.append(text)

    if para_texts:
        blocks.append({
            "text": "\n".join(para_texts),
            "source": {
                "source_type": "docx",
                "page_number": None,
                "sheet_name": None,
                "row_start": None,
                "row_end": None,
            }
        })

    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))

        if rows:
            blocks.append({
                "text": "\n".join(rows),
                "source": {
                    "source_type": "docx_table",
                    "page_number": None,
                    "sheet_name": f"table_{table_index + 1}",
                    "row_start": None,
                    "row_end": None,
                }
            })

    return blocks